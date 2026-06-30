from pathlib import Path

from docx import Document as DocxDocument

from greatocr.model.document import Document, Issue, Page
from greatocr.reports.quality_docx import write_quality_docx
from greatocr.validation.quality import compute_quality_summary


def make_summary(issues: list[Issue]):
    document = Document(
        document_id="doc-1",
        source_file_name="sample.pdf",
        file_sha256="a" * 64,
        page_count=1,
        provider_name="fake",
        pages=[
            Page(
                page_id="page-0001",
                page_number=1,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
            )
        ],
    )
    return compute_quality_summary(document, issues)


def make_issue() -> Issue:
    return Issue(
        issue_id="issue-1",
        page_number=1,
        issue_type="critical_field_low_confidence",
        severity="high",
        message="Amount confidence is low.",
        snippet="USD 1,2?0.00",
        suggestion="Review manually.",
    )


def docx_text(path: Path) -> str:
    document = DocxDocument(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_quality_docx_title_contains_chinese_report_name(tmp_path: Path) -> None:
    output = write_quality_docx(make_summary([]), [], tmp_path / "quality-report.docx")

    assert "GreatOCR 质量报告" in docx_text(output)


def test_quality_docx_contains_issue_details(tmp_path: Path) -> None:
    issue = make_issue()
    output = write_quality_docx(make_summary([issue]), [issue], tmp_path / "quality-report.docx")
    text = docx_text(output)

    assert "页码：1" in text
    assert "原文片段：USD 1,2?0.00" in text
    assert "问题类型：critical_field_low_confidence" in text
    assert "建议：Review manually." in text


def test_quality_docx_without_issues_says_no_key_risks(tmp_path: Path) -> None:
    output = write_quality_docx(make_summary([]), [], tmp_path / "quality-report.docx")

    assert "未发现关键风险" in docx_text(output)
