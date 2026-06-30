from pathlib import Path
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.shared import Pt

from greatocr.docx.builder import build_docx
from greatocr.model.document import Block, Document, Page, TextSpan


def text_block(block_type: str, order: int, text: str) -> Block:
    return Block(
        block_id=f"block-{order}",
        block_type=block_type,
        reading_order=order,
        spans=[TextSpan(span_id=f"span-{order}", original_text=text, current_text=text)],
    )


def make_document(pages: list[Page]) -> Document:
    return Document(
        document_id="doc-1",
        source_file_name="sample.pdf",
        file_sha256="a" * 64,
        page_count=len(pages),
        provider_name="fake",
        pages=pages,
    )


def test_minimal_document_generates_docx_file(tmp_path: Path) -> None:
    document = make_document(
        [
            Page(
                page_id="page-0001",
                page_number=1,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
                blocks=[text_block("paragraph", 1, "Hello")],
            )
        ]
    )

    result = build_docx(document, tmp_path / "result.docx")

    assert result.output_path.is_file()
    with ZipFile(result.output_path) as package:
        assert "[Content_Types].xml" in package.namelist()
        assert "word/document.xml" in package.namelist()


def test_title_and_paragraph_are_readable_with_python_docx(tmp_path: Path) -> None:
    document = make_document(
        [
            Page(
                page_id="page-0001",
                page_number=1,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
                blocks=[
                    text_block("title", 1, "Sample Contract"),
                    text_block("paragraph", 2, "Editable paragraph."),
                ],
            )
        ]
    )

    result = build_docx(document, tmp_path / "result.docx")
    docx = DocxDocument(result.output_path)

    assert [p.text for p in docx.paragraphs[:2]] == [
        "Sample Contract",
        "Editable paragraph.",
    ]


def test_multi_page_document_uses_sections_and_reports_pagination_risk(tmp_path: Path) -> None:
    document = make_document(
        [
            Page(
                page_id="page-0001",
                page_number=1,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
                blocks=[text_block("paragraph", 1, "Page one")],
            ),
            Page(
                page_id="page-0002",
                page_number=2,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
                blocks=[text_block("paragraph", 1, "Page two")],
            ),
        ]
    )

    result = build_docx(document, tmp_path / "result.docx")

    reopened = DocxDocument(result.output_path)
    assert len(reopened.sections) == 2
    assert result.issues[0].issue_type == "pagination_may_drift"


def test_page_size_is_written_to_section(tmp_path: Path) -> None:
    document = make_document(
        [
            Page(
                page_id="page-0001",
                page_number=1,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
                blocks=[text_block("paragraph", 1, "Sized page")],
            )
        ]
    )

    result = build_docx(document, tmp_path / "result.docx")
    section = DocxDocument(result.output_path).sections[0]

    assert section.page_width == Pt(612)
    assert section.page_height == Pt(792)


def test_rotated_source_generates_portrait_section(tmp_path: Path) -> None:
    page = Page(
        page_id="page-0001",
        page_number=1,
        width=842,
        height=595,
        rotation=270,
        page_type="scanned",
        blocks=[text_block("paragraph", 1, "Portrait page")],
    )

    result = build_docx(make_document([page]), tmp_path / "portrait.docx")
    section = DocxDocument(result.output_path).sections[0]

    assert section.page_height > section.page_width
    assert section.page_width == Pt(595)
    assert section.page_height == Pt(842)


def test_header_and_footer_use_word_parts(tmp_path: Path) -> None:
    page = Page(
        page_id="page-0001",
        page_number=1,
        width=612,
        height=792,
        rotation=0,
        page_type="native_text",
        blocks=[
            text_block("header", 1, "Company Header"),
            text_block("paragraph", 2, "Body text"),
            text_block("footer", 3, "Page Footer"),
        ],
    )

    result = build_docx(make_document([page]), tmp_path / "header-footer.docx")
    reopened = DocxDocument(result.output_path)

    assert any("Company Header" in p.text for p in reopened.sections[0].header.paragraphs)
    assert any("Page Footer" in p.text for p in reopened.sections[0].footer.paragraphs)
    assert [p.text for p in reopened.paragraphs] == ["Body text"]
