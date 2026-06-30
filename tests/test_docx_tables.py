from pathlib import Path
from zipfile import ZipFile

from docx import Document as DocxDocument

from greatocr.docx.builder import build_docx
from greatocr.model.document import Block, Document, Page, Table, TableCell


def make_document(table: Table) -> Document:
    return Document(
        document_id="doc-1",
        source_file_name="sample.pdf",
        file_sha256="a" * 64,
        page_count=1,
        provider_name="fake",
        pages=[
            Page(
                page_id="page-0001",
                page_number=1,
                width=612,
                height=792,
                rotation=0,
                page_type="native_text",
                blocks=[
                    Block(
                        block_id="block-table",
                        block_type="table",
                        reading_order=1,
                        table=table,
                    )
                ],
            )
        ],
    )


def test_basic_table_generates_real_word_table(tmp_path: Path) -> None:
    table = Table(
        table_id="table-1",
        rows=[
            [TableCell(text="Item"), TableCell(text="Amount")],
            [TableCell(text="Service"), TableCell(text="100.00")],
        ],
    )

    result = build_docx(make_document(table), tmp_path / "result.docx")
    docx = DocxDocument(result.output_path)

    assert len(docx.tables) == 1
    assert docx.tables[0].cell(1, 1).text == "100.00"


def test_colspan_cell_writes_word_merge(tmp_path: Path) -> None:
    table = Table(
        table_id="table-1",
        rows=[
            [TableCell(text="Heading", col_span=2), TableCell(text="")],
            [TableCell(text="Service"), TableCell(text="100.00")],
        ],
    )

    result = build_docx(make_document(table), tmp_path / "result.docx")
    xml = ZipFile(result.output_path).read("word/document.xml").decode("utf-8")

    assert "w:gridSpan" in xml


def test_low_confidence_table_is_degraded_to_issue(tmp_path: Path) -> None:
    table = Table(
        table_id="table-1",
        rows=[[TableCell(text="Uncertain")]],
        confidence=0.4,
    )

    result = build_docx(make_document(table), tmp_path / "result.docx")
    docx = DocxDocument(result.output_path)

    assert len(docx.tables) == 0
    assert result.issues[0].issue_type == "table_degraded"
