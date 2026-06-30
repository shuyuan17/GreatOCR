from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument

from greatocr.model.document import Issue
from greatocr.validation.quality import QualitySummary


def write_quality_docx(
    summary: QualitySummary,
    issues: list[Issue],
    output_path: Path,
) -> Path:
    document = WordDocument()
    document.add_heading("GreatOCR 质量报告", level=1)
    document.add_paragraph(f"文件名：{summary.file_name}")
    document.add_paragraph(f"页数：{summary.page_count}")
    document.add_paragraph(f"解析供应商：{summary.provider_name}")
    document.add_paragraph(f"总体质量评级：{summary.rating}")
    document.add_paragraph(f"页面类型统计：{summary.page_type_counts}")
    document.add_paragraph(f"表格降级数量：{summary.table_degraded_count}")
    document.add_paragraph(f"字体替换数量：{summary.font_substitution_count}")
    document.add_paragraph(f"自动校正数量：{summary.auto_correction_count}")

    document.add_heading("问题列表", level=2)
    if not issues:
        document.add_paragraph("未发现关键风险")
    else:
        for item in issues:
            document.add_paragraph(f"页码：{item.page_number}")
            document.add_paragraph(f"原文片段：{item.snippet or ''}")
            document.add_paragraph(f"问题类型：{item.issue_type}")
            document.add_paragraph(f"说明：{item.message}")
            document.add_paragraph(f"建议：{item.suggestion or '请人工复核。'}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
